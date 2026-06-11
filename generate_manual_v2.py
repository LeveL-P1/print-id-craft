"""
WiseMelon User Manual Generator - With blank screenshot placeholders
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "WiseMelon_User_Manual.docx")

def h(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for r in heading.runs:
        r.font.color.rgb = RGBColor(0x18, 0x18, 0x37)

def placeholder(doc, label, height_inches=3.5):
    """Add a labeled blank box for screenshot placement"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add vertical space
    for _ in range(int(height_inches * 4)):
        p.add_run("\n")
    run = p.add_run(f"📸 {label}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.italic = True
    p.add_run("\n\n")
    run2 = p.add_run("[Paste screenshot here]")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    doc.add_paragraph()

def step(doc, num, text):
    p = doc.add_paragraph()
    r = p.add_run(f"Step {num}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    p.add_run(text)

def info_box(doc, title, content):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EFF6FF"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    r = p.add_run(f"💡 {title}: ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    r2 = p.add_run(content)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    doc.add_paragraph()

def warn_box(doc, title, content):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FEF2F2"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    r = p.add_run(f"⚠️ {title}: ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    r2 = p.add_run(content)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    doc.add_paragraph()

def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ══════════ TITLE PAGE ══════════
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_heading('WiseMelon', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.size = Pt(36)
        r.font.color.rgb = RGBColor(0x18, 0x18, 0x37)
    
    sub = doc.add_paragraph('Comprehensive User Manual')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(18)
    sub.runs[0].font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    
    doc.add_paragraph()
    desc = doc.add_paragraph('Multi-School ID Card Management & Print Portal')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.runs[0].font.size = Pt(14)
    desc.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    for _ in range(5):
        doc.add_paragraph()
    ver = doc.add_paragraph('Version 1.0  •  April 2026')
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    doc.add_page_break()

    # ══════════ TABLE OF CONTENTS ══════════
    h(doc, 'Table of Contents')
    toc = [
        "1. Introduction & Overview",
        "2. System Architecture & User Roles",
        "3. Landing Page",
        "4. Login System",
        "   4.1 Teacher Login",
        "   4.2 Manufacturer (Admin) Login",
        "5. Manufacturer Dashboard",
        "   5.1 Dashboard Overview & Statistics",
        "   5.2 Recent Schools Table",
        "6. Schools Management",
        "   6.1 Schools List & Cards",
        "   6.2 Adding a New School",
        "   6.3 Search, Filter & Pagination",
        "   6.4 Deleting a School",
        "7. School Detail Page",
        "   7.1 Overview Tab & Teacher Credentials",
        "   7.2 School Logo Upload",
        "   7.3 Classes Tab — Form Links & Sharing",
        "   7.4 Students Tab — Management & Actions",
        "   7.5 Template Tab — JPG Template Mapper",
        "   7.6 Generate Tab — Batch Generation",
        "   7.7 Batches Tab — Downloads",
        "   7.8 Export Tab — CSV & Excel",
        "   7.9 Bulk Import & Bulk Photo Upload",
        "8. Teacher Dashboard",
        "   8.1 Overview Tab & Stats",
        "   8.2 Class Form Links & Sharing",
        "   8.3 Class Breakdown & Export",
        "   8.4 Students Tab — Review & Actions",
        "   8.5 Sub-Teachers Management",
        "   8.6 ID Template Preview",
        "9. Student Submission Form",
        "   9.1 Step 1: Fill Details",
        "   9.2 Step 2: Photo Upload & Verification",
        "   9.3 Step 3: Review & Submit",
        "   9.4 Success Screen",
        "10. Photo Verification Engine",
        "11. ID Card Preview & Watermark System",
        "12. Data Export (CSV & Excel)",
        "13. Security Features",
        "14. Troubleshooting & FAQ",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        if not item.startswith("   "):
            p.runs[0].bold = True
        else:
            p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_page_break()

    # ══════════ 1. INTRODUCTION ══════════
    h(doc, '1. Introduction & Overview')
    doc.add_paragraph(
        'WiseMelon is a comprehensive, web-based platform designed for managing student ID cards '
        'across multiple schools. It provides a complete end-to-end solution from student data collection '
        'to ID card printing, with built-in role-based access control, photo verification, and batch processing.')
    doc.add_paragraph()
    h(doc, 'Key Features', 2)
    features = [
        'Multi-School Management — Manage unlimited schools from a single manufacturer dashboard',
        'Role-Based Access — Three distinct roles: Manufacturer, Main Teacher, Sub-Teacher',
        'Smart Form Links — Unique per-class submission URLs for students/parents',
        'Photo Verification Engine — Automated quality checks for passport-size photos',
        'JPG Template Mapping — Upload ID card images and visually map data fields',
        'Dual-Side Support — Front and back ID card template support',
        'Batch Generation — Print-ready PDF batches with all approved student cards',
        'Data Export — CSV and Excel export with class/status filters',
        'Bulk Import — Import student data from CSV/Excel with validation preview',
        'Bulk Photo Upload — Folder upload with automatic name-matching to students',
        'Sub-Teacher System — Class-specific teacher accounts with restricted visibility',
        'Watermarked Previews — Mandatory watermarks on all card previews',
    ]
    bullets(doc, features)
    doc.add_page_break()

    # ══════════ 2. ARCHITECTURE ══════════
    h(doc, '2. System Architecture & User Roles')
    doc.add_paragraph(
        'WiseMelon uses Next.js, Prisma ORM, PostgreSQL (Supabase), and Supabase Storage. '
        'The application supports three distinct user roles:')
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    for i, hdr in enumerate(['Role', 'Access Level', 'Key Permissions', 'Login URL']):
        table.cell(0, i).text = hdr
        for r in table.cell(0, i).paragraphs[0].runs: r.bold = True
    
    roles = [
        ('Manufacturer (Admin)', 'Full System', 'Create schools, manage templates, view all data, generate batches, bulk download cards', '/login?mode=admin'),
        ('Main Teacher', 'School-wide', 'View all classes, approve/flag students, add sub-teachers, add classes, export data, map templates', '/login'),
        ('Sub-Teacher (Class)', 'Single Class', 'View assigned class only, approve/flag own students, add comments', '/login'),
    ]
    for idx, (role, access, perms, url) in enumerate(roles, 1):
        table.cell(idx, 0).text = role
        table.cell(idx, 1).text = access
        table.cell(idx, 2).text = perms
        table.cell(idx, 3).text = url
    
    doc.add_paragraph()
    info_box(doc, "Workflow", "Manufacturer creates school → Assigns Main Teacher → Main Teacher adds classes & sub-teachers → Form links shared with students → Students submit data & photos → Teachers review & approve → Manufacturer generates print batches → Cards printed & delivered.")
    doc.add_page_break()

    # ══════════ 3. LANDING PAGE ══════════
    h(doc, '3. Landing Page')
    doc.add_paragraph('The landing page introduces WiseMelon with animated sections and a modern design.')
    
    doc.add_paragraph()
    h(doc, '3.1 Hero Section', 2)
    doc.add_paragraph('Features the "We build powerful ID Card systems" tagline, CTA buttons (Explore Solution, School Login), and information cards highlighting platform differentiators on the right side.')
    placeholder(doc, "Figure 3.1: Landing Page — Hero Section")
    
    h(doc, '3.2 Statistics Bar', 2)
    doc.add_paragraph('Animated statistics bar with: 50+ Schools Managed, 25,000+ Cards Printed, 99.9% Match Accuracy. Numbers animate from zero as user scrolls into view.')
    placeholder(doc, "Figure 3.2: Statistics Bar", 2.5)
    
    h(doc, '3.3 Partnership Section', 2)
    doc.add_paragraph('Addresses common concerns about ID card management with a "Partnership-First Approach" message. Includes an animated submarine SVG illustration.')
    placeholder(doc, "Figure 3.3: Partnership Section")
    
    h(doc, '3.4 Footer & CTA', 2)
    doc.add_paragraph('Dark-themed footer with "Get Started" CTA button, links to Privacy Policy, Terms of Service, and Support.')
    placeholder(doc, "Figure 3.4: Footer Section", 2.5)
    doc.add_page_break()

    # ══════════ 4. LOGIN SYSTEM ══════════
    h(doc, '4. Login System')
    doc.add_paragraph('Separate login portals for Teachers and Manufacturers using NextAuth.js with bcrypt password hashing.')
    
    doc.add_paragraph()
    h(doc, '4.1 Teacher Login', 2)
    doc.add_paragraph('Located at /login. Split-screen layout with animated ID card mock-up on the left, login form on the right. Photo guidelines reminder displayed below the form.')
    placeholder(doc, "Figure 4.1: Teacher Login Page")
    bullets(doc, [
        'Email and password authentication',
        'Photo guidelines reminder for student submissions',
        'Animated ID card illustration with floating dots',
        'Responsive design for mobile devices',
        'Toast notifications for success/error feedback',
    ])
    
    doc.add_paragraph()
    h(doc, '4.2 Manufacturer (Admin) Login', 2)
    doc.add_paragraph('Located at /login?mode=admin. Distinct "Manufacturer Login" header with "Manage schools, templates & printing" subtitle and factory icon.')
    placeholder(doc, "Figure 4.2: Manufacturer Login Page")
    
    doc.add_paragraph()
    doc.add_paragraph('Login Steps:')
    step(doc, 1, 'Navigate to the login page (/login for teachers, /login?mode=admin for manufacturer)')
    step(doc, 2, 'Enter your registered email address')
    step(doc, 3, 'Enter your password')
    step(doc, 4, 'Click "Sign In" — you will be redirected to the appropriate dashboard')
    
    info_box(doc, "Default Credentials", "Main teacher accounts use default password: Teacher@123. The manufacturer should share these credentials with the school administrator.")
    doc.add_page_break()

    # ══════════ 5. MANUFACTURER DASHBOARD ══════════
    h(doc, '5. Manufacturer Dashboard')
    doc.add_paragraph('The central command center for managing all schools, templates, and printing operations. Includes a left sidebar with navigation.')
    
    h(doc, '5.1 Dashboard Overview & Statistics', 2)
    doc.add_paragraph('Four animated statistics cards:')
    bullets(doc, [
        'Total Schools (Blue) — Number of registered schools',
        'Total Students (Purple) — Total student submissions across all schools',
        'Print Batches (Amber) — Total generated print batches',
        'This Month (Green) — Student submissions for the current month',
    ])
    placeholder(doc, "Figure 5.1: Manufacturer Dashboard with Stats Cards")
    
    h(doc, '5.2 Recent Schools Table', 2)
    doc.add_paragraph('Displays 5 most recently added schools with columns: School name & email, Student count, Classes count, Template status (Ready/Not Set), and Manage button. On mobile, transforms to stacked cards.')
    placeholder(doc, "Figure 5.2: Recent Schools Table", 2.5)
    
    h(doc, 'Sidebar Navigation', 3)
    bullets(doc, [
        'Dashboard — Overview with stats and recent schools',
        'Schools — Complete school list with management',
        'User Profile — Logged-in user name with sign-out',
    ])
    doc.add_page_break()

    # ══════════ 6. SCHOOLS MANAGEMENT ══════════
    h(doc, '6. Schools Management')
    
    h(doc, '6.1 Schools List & Cards', 2)
    doc.add_paragraph('Schools displayed as gradient-banner cards showing: school initial, name, contact, student/class counts, template status badge, Manage and Delete buttons. Summary stats at top.')
    placeholder(doc, "Figure 6.1: Schools List with Cards")
    
    h(doc, '6.2 Adding a New School', 2)
    doc.add_paragraph('Click "+ Add School" to open the modal:')
    step(doc, 1, 'Fill School Name (required) and Contact Email (required)')
    step(doc, 2, 'Optionally add Address and upload School Logo')
    step(doc, 3, 'Add class names (e.g., "1st A", "2nd B") — press Enter or click + Add')
    step(doc, 4, 'Click "Create School" — auto-generates main teacher account and class form links')
    placeholder(doc, "Figure 6.2: Add School Modal")
    
    h(doc, '6.3 Search, Filter & Pagination', 2)
    doc.add_paragraph('Debounced search bar (400ms) for name, email, or address. Clear button resets filter. Pagination for 50+ schools per page.')
    
    h(doc, '6.4 Deleting a School', 2)
    doc.add_paragraph('Click red trash icon → type "DELETE" to confirm. Permanently removes school and ALL associated data.')
    warn_box(doc, "Warning", "Deleting a school is irreversible. All student data, photos, classes, templates, and batches will be permanently lost.")
    doc.add_page_break()

    # ══════════ 7. SCHOOL DETAIL ══════════
    h(doc, '7. School Detail Page')
    doc.add_paragraph('7 tabs: Overview, Classes, Students, Template, Generate, Batches, Export. Breadcrumb navigation: Dashboard › Schools › School Name.')
    
    h(doc, '7.1 Overview Tab & Teacher Credentials', 2)
    doc.add_paragraph('Shows 4 stat cards (Classes, Students, Batches, Template Status). Below, the Main Teacher Login section displays:')
    bullets(doc, [
        'Login URL with Copy button',
        'Email address with Copy button',
        'Default password: Teacher@123',
        'Reset Password button (restores default)',
        'Quick Auto-Generate option if no teacher exists',
        'Manual email entry option',
    ])
    placeholder(doc, "Figure 7.1: School Overview with Teacher Credentials")
    
    h(doc, '7.2 School Logo Upload', 2)
    doc.add_paragraph('Upload/replace school logo (JPEG, PNG, WebP). Drag-and-drop supported. Logo appears on student ID cards.')
    placeholder(doc, "Figure 7.2: School Logo Upload Section", 2)
    
    h(doc, '7.3 Classes Tab — Form Links & Sharing', 2)
    doc.add_paragraph('Complete class management:')
    bullets(doc, [
        'Add New Class — name + optional expiry date',
        'Unique Form Link per class — students use this to submit details',
        'Copy Link — one-click clipboard copy',
        'Share via WhatsApp — pre-formatted message with school/class info',
        'Share via Email — pre-filled subject and body',
        'Toggle Active/Inactive — enable/disable submissions',
        'Set Expiry Date — auto-deactivate after date',
        'Delete Class — requires "DELETE" confirmation',
    ])
    placeholder(doc, "Figure 7.3: Classes Tab with Form Links")
    
    h(doc, '7.4 Students Tab — Management & Actions', 2)
    doc.add_paragraph('Student table with filters (Status, Class, Search) and columns: Photo, Serial, Name, Class, Status, Flag Note, Teacher Comment, Actions.')
    doc.add_paragraph('Available actions per student:')
    bullets(doc, [
        'View (👁) — Full detail modal with photo, form data, ID card preview',
        'Approve — Mark approved for printing',
        'Flag (🚩) — Flag with custom reason note',
        'Unflag — Remove flag, return to submitted',
        'Edit — Modify form fields in modal',
        'Delete — Remove student record',
    ])
    placeholder(doc, "Figure 7.4: Students Tab with Table")
    
    h(doc, '7.5 Template Tab — JPG Template Mapper', 2)
    doc.add_paragraph('Upload an actual ID card template image and visually map data field positions:')
    step(doc, 1, 'Upload front-side template image (JPG/PNG)')
    step(doc, 2, 'Optionally upload back-side template')
    step(doc, 3, 'Click on template to add field mapping zones')
    step(doc, 4, 'For each zone: select data field (Full Name, Class, Photo, Serial, etc.)')
    step(doc, 5, 'Adjust font size, color, alignment, position')
    step(doc, 6, 'Preview with sample data; Save configuration')
    placeholder(doc, "Figure 7.5: JPG Template Mapper")
    
    h(doc, '7.6 Generate Tab — Batch Generation', 2)
    doc.add_paragraph('Preview cards with actual student data, then generate print-ready batches:')
    step(doc, 1, 'Review card preview with real student photos and data')
    step(doc, 2, 'Click "Generate Batch" to start processing')
    step(doc, 3, 'Progress indicator shows status; auto-polls for completion')
    step(doc, 4, 'Completed batch appears in Batches tab for download')
    info_box(doc, "Note", "Only SUBMITTED or APPROVED students are included. Flagged students are excluded. High-quality downloads restricted to Manufacturer role.")
    placeholder(doc, "Figure 7.6: Generate Tab with Preview")
    
    h(doc, '7.7 Batches Tab — Downloads', 2)
    doc.add_paragraph('Lists all generated batches with: creation date, student count, status (Processing/Ready), download options (Front PDF, Back PDF, Manifest CSV).')
    placeholder(doc, "Figure 7.7: Batches Tab", 2)
    
    h(doc, '7.8 Export Tab — CSV & Excel', 2)
    doc.add_paragraph('Download student data filtered by class and status. CSV (.csv) and Excel (.xlsx) formats. Includes: Serial Number, Name, Class, all form fields, Photo URL, Status, Comments, Date.')
    placeholder(doc, "Figure 7.8: Export Tab", 2)
    
    h(doc, '7.9 Bulk Import & Bulk Photo Upload', 2)
    doc.add_paragraph('Bulk Student Import:')
    step(doc, 1, 'Click "📥 Import" → select class → upload CSV/Excel')
    step(doc, 2, 'Review validation preview (matched columns, errors, warnings)')
    step(doc, 3, 'Confirm to import data into system')
    doc.add_paragraph('Bulk Photo Upload:')
    step(doc, 1, 'Click "📷 Bulk Photos" → select folder with student photos')
    step(doc, 2, 'Auto-matches by filename (serial number or student name)')
    step(doc, 3, 'Review results; manually assign unmatched photos')
    placeholder(doc, "Figure 7.9: Bulk Import / Photo Upload", 2.5)
    doc.add_page_break()

    # ══════════ 8. TEACHER DASHBOARD ══════════
    h(doc, '8. Teacher Dashboard')
    doc.add_paragraph('Central hub for school administrators and class teachers. Tabs: Overview, Students, Sub-Teachers (main only), ID Template (main only).')
    
    h(doc, '8.1 Overview Tab & Stats', 2)
    doc.add_paragraph('Displays school name, teacher role badge (Main Teacher / Class Teacher), and stat cards:')
    bullets(doc, [
        'Total Submissions — all student entries',
        'Approved — approved for printing (green)',
        'Pending Review — awaiting teacher review (amber)',
        'Flagged — marked for issues (red)',
        'Approval Progress bar with percentage',
    ])
    placeholder(doc, "Figure 8.1: Teacher Dashboard Overview")
    
    h(doc, '8.2 Class Form Links & Sharing', 2)
    doc.add_paragraph('Each class shows a unique submission URL with: 📋 Copy button, 💬 WhatsApp share (pre-formatted message), student count, assigned teacher name. Main teachers can add new classes via "+ Add Class" form.')
    placeholder(doc, "Figure 8.2: Class Form Links Section")
    
    h(doc, '8.3 Class Breakdown & Export', 2)
    doc.add_paragraph('Table with columns: Class, Teacher, Students, Approved, Flagged. Export buttons for CSV and Excel formats.')
    placeholder(doc, "Figure 8.3: Class Breakdown Table", 2)
    
    h(doc, '8.4 Students Tab — Review & Actions', 2)
    doc.add_paragraph('Filters: Class dropdown (main teachers only), Status dropdown. Student table columns: Photo, Serial, Name, Class, Status, Comment, Actions.')
    doc.add_paragraph('Action buttons per student:')
    bullets(doc, [
        '👁 View — Detailed modal with photo, form data, ID card preview',
        '✓ Approve — Mark approved for printing',
        '✕ Disapprove — Mark as flagged/disapproved',
        '🚩 Flag — Flag with custom reason note',
        'Unflag — Remove flag',
        '✏️ Edit — Modify all form fields',
        '💬 Comment — Add teacher comment (visible to manufacturer)',
    ])
    placeholder(doc, "Figure 8.4: Teacher Students Tab")
    
    h(doc, '8.5 Sub-Teachers Management', 2)
    doc.add_paragraph('Available only to Main Teachers. Add class-specific teacher accounts:')
    step(doc, 1, 'Click "+ Add Teacher"')
    step(doc, 2, 'Fill: Name, Email, Password (min 6 chars), Assign to Class')
    step(doc, 3, 'Success banner shows credentials (Name, Email, Password, Login URL)')
    step(doc, 4, 'Copy & share — password shown ONLY once')
    doc.add_paragraph('Sub-teacher list table: Name (with avatar), Email, Assigned Class, Added date, Remove button. Sub-teachers can only see their assigned class.')
    placeholder(doc, "Figure 8.5: Sub-Teachers Management")
    warn_box(doc, "Security", "Passwords are bcrypt-hashed and shown only at creation. If forgotten, delete and recreate the sub-teacher account.")
    
    h(doc, '8.6 ID Template Preview', 2)
    doc.add_paragraph('Main teachers can view and configure the school\'s ID card template using the JPG Template Mapper. All previews include mandatory watermarks.')
    placeholder(doc, "Figure 8.6: Teacher Template Preview", 2)
    doc.add_page_break()

    # ══════════ 9. STUDENT FORM ══════════
    h(doc, '9. Student Submission Form')
    doc.add_paragraph('Students/parents use a unique class URL (/submit/{token}) to submit ID card data. 3-step wizard: Details → Photo → Review.')
    
    h(doc, '9.1 Step 1: Fill Details', 2)
    doc.add_paragraph('Displays school name, class, step indicator (1-2-3). Form fields include: Full Name, Father\'s Name, DOB, Blood Group, Address, etc. Required fields marked with *.')
    placeholder(doc, "Figure 9.1: Student Form — Details Step")
    
    h(doc, '9.2 Step 2: Photo Upload & Verification', 2)
    doc.add_paragraph('Photo requirements box lists: 300px minimum, plain background, face visible, 3:4 ratio. Includes sample photo reference. Photo Verifier checks quality before accepting. React Image Crop for passport-size cropping.')
    placeholder(doc, "Figure 9.2: Student Form — Photo Upload & Crop")
    
    h(doc, '9.3 Step 3: Review & Submit', 2)
    doc.add_paragraph('Side-by-side: Card Preview (with watermark) and Details Check panel. Shows student photo thumbnail, all entered fields, class, and confirmation note. Progress bar during upload (0-100%).')
    placeholder(doc, "Figure 9.3: Student Form — Review Screen")
    
    h(doc, '9.4 Success Screen', 2)
    p = doc.add_paragraph('Green checkmark, unique Serial Number (students should save this), watermarked ID Card Preview, "Powered by WiseMelon" branding.')
    placeholder(doc, "Figure 9.4: Submission Success Screen", 2)
    doc.add_page_break()

    # ══════════ 10. PHOTO VERIFICATION ══════════
    h(doc, '10. Photo Verification Engine')
    doc.add_paragraph('Built-in automated photo validation:')
    checks = [
        '🔍 Minimum Resolution — 300×300 pixels for print quality',
        '🔍 File Format — JPEG, PNG, WebP only',
        '🔍 File Size — Maximum 5MB',
        '🔍 Aspect Ratio — Cropped to 3:4 (passport standard)',
        '🔍 Background Check — Verifies plain/solid background',
        '🔍 Face Detection — Ensures face is visible and centered',
    ]
    bullets(doc, checks)
    info_box(doc, "Photo Tips", "Use natural lighting, plain white/light wall, face camera directly, avoid filters or editing.")
    placeholder(doc, "Figure 10: Photo Verification in Action", 2.5)

    # ══════════ 11. WATERMARK ══════════
    h(doc, '11. ID Card Preview & Watermark System')
    doc.add_paragraph('Two preview types: JPG Card Preview (template image with data overlay) and Canvas-Based Preview (drag-and-drop layout).')
    doc.add_paragraph('Watermark system:')
    bullets(doc, [
        'Diagonal watermark on ALL preview screens (student, teacher, manufacturer)',
        'Removed only in high-quality batch downloads (Manufacturer role only)',
        'Cannot be disabled by teachers or students',
        'Covers entire card surface to prevent screenshot reproduction',
    ])
    placeholder(doc, "Figure 11: ID Card Preview with Watermark", 2.5)
    doc.add_page_break()

    # ══════════ 12. EXPORT ══════════
    h(doc, '12. Data Export (CSV & Excel)')
    doc.add_paragraph('Both manufacturer and teacher dashboards support export. Includes: Serial Number, Name, Class, all form fields, Photo URL, Status, Flag Note, Teacher Comment, Date.')
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    for i, hdr in enumerate(['Feature', 'CSV', 'Excel']):
        table.cell(0, i).text = hdr
        for r in table.cell(0, i).paragraphs[0].runs: r.bold = True
    table.cell(1, 0).text = 'Format'; table.cell(1, 1).text = '.csv'; table.cell(1, 2).text = '.xlsx formatted'
    table.cell(2, 0).text = 'Filters'; table.cell(2, 1).text = 'Class, Status'; table.cell(2, 2).text = 'Class, Status'
    doc.add_paragraph()

    # ══════════ 13. SECURITY ══════════
    h(doc, '13. Security Features')
    security = [
        '🔒 Authentication — NextAuth.js with JWT sessions (24h expiry), bcrypt passwords',
        '🔒 Role-Based Access — Middleware enforces role-based routing',
        '🔒 Sub-Teacher Isolation — Class teachers see only their assigned class',
        '🔒 Watermarked Previews — Mandatory on all card previews',
        '🔒 Batch Download Restriction — High-quality downloads restricted to Manufacturer',
        '🔒 Form Link Control — Deactivation and expiry date support',
        '🔒 Delete Confirmation — Requires typing "DELETE" for destructive operations',
        '🔒 Input Validation — Client and server-side validation',
        '🔒 Photo Validation — Format, size, resolution, and content quality checks',
        '🔒 Force Logout — API endpoint to force sign-out user sessions',
    ]
    bullets(doc, security)
    doc.add_page_break()

    # ══════════ 14. FAQ ══════════
    h(doc, '14. Troubleshooting & FAQ')
    faqs = [
        ('Q: I forgot my teacher login password.',
         'A: Contact the manufacturer. They can reset from School Detail → Overview → "Reset Pw" button. Default: Teacher@123.'),
        ('Q: A student submitted incorrect information.',
         'A: Students tab → Edit (✏️) button to modify form fields. Or flag the student and add a comment.'),
        ('Q: The form link isn\'t working.',
         'A: Check if class link is active (Classes tab → Toggle). Check expiry date. Generate new link if needed.'),
        ('Q: Photos are being rejected.',
         'A: Ensure: 300×300px minimum, JPEG/PNG/WebP, under 5MB, plain background, front-facing.'),
        ('Q: How do I generate print-ready cards?',
         'A: School → Generate tab → Review preview → Click "Generate Batch". Download from Batches tab.'),
        ('Q: Can sub-teachers see all students?',
         'A: No. Sub-teachers only see their assigned class. Only Main Teacher has school-wide access.'),
        ('Q: How to share form links with parents?',
         'A: Teacher Dashboard → Overview → Class Form Links → Copy or WhatsApp share button.'),
        ('Q: Can I import from spreadsheet?',
         'A: Yes. Students tab → "📥 Import" → upload CSV/Excel → validate → confirm import.'),
        ('Q: Why are previews watermarked?',
         'A: Watermarks prevent unauthorized reproduction. Removed only in Manufacturer batch downloads.'),
        ('Q: How to upload photos in bulk?',
         'A: Students tab → "📷 Bulk Photos" → select folder → auto-matches by filename → manually assign unmatched.'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        r = p.add_run(q)
        r.bold = True
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        doc.add_paragraph(a)
        doc.add_paragraph()

    # ══════════ END PAGE ══════════
    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph()
    end = doc.add_heading('Thank You', 1)
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('For support or feature requests, contact the WiseMelon team.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p2 = doc.add_paragraph('© 2026 WiseMelon. All rights reserved.')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(OUTPUT)
    print(f"✅ Manual saved: {OUTPUT}")

if __name__ == "__main__":
    main()

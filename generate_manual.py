"""
Generate a comprehensive User Manual for Print ID Craft in Word format.
Uses actual screenshots captured from the application.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCREENSHOTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "screenshots")
OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Print_ID_Craft_User_Manual.docx")

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x18, 0x18, 0x37)
    return heading

def add_screenshot(doc, filename, caption="", width=6.0):
    """Add a screenshot image with caption."""
    path = os.path.join(SCREENSHOTS_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.style.font.size = Pt(9)
            cap.style.font.italic = True
            cap.style.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        return True
    else:
        p = doc.add_paragraph(f"[Screenshot: {filename} - Not available]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        return False

def add_info_box(doc, title, content, box_type="info"):
    """Add a styled info/tip/warning box."""
    colors = {
        "info": ("EFF6FF", "3B82F6", "1E40AF"),
        "tip": ("F0FDF4", "22C55E", "15803D"),
        "warning": ("FEF2F2", "EF4444", "B91C1C"),
        "note": ("F8FAFC", "64748B", "334155"),
    }
    bg, border, text_color = colors.get(box_type, colors["info"])
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg)
    
    p = cell.paragraphs[0]
    run = p.add_run(f"💡 {title}: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(int(text_color[:2], 16), int(text_color[2:4], 16), int(text_color[4:], 16))
    
    run2 = p.add_run(content)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(int(text_color[:2], 16), int(text_color[2:4], 16), int(text_color[4:], 16))
    
    doc.add_paragraph()

def add_step(doc, number, text):
    """Add a numbered step."""
    p = doc.add_paragraph()
    run = p.add_run(f"Step {number}: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)

def create_manual():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1B, 0x1C, 0x1C)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ═══════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_heading('Print ID Craft', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0x18, 0x18, 0x37)
    
    subtitle = doc.add_paragraph('Comprehensive User Manual')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    
    doc.add_paragraph()
    
    desc = doc.add_paragraph('Multi-School ID Card Management & Print Portal')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.runs[0].font.size = Pt(14)
    desc.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    for _ in range(4):
        doc.add_paragraph()
    
    version = doc.add_paragraph('Version 1.0  •  April 2026')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.runs[0].font.size = Pt(11)
    version.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, 'Table of Contents', level=1)
    
    toc_items = [
        ("1.", "Introduction & Overview"),
        ("2.", "System Architecture & User Roles"),
        ("3.", "Landing Page"),
        ("4.", "Login System"),
        ("  4.1", "Teacher Login"),
        ("  4.2", "Manufacturer (Admin) Login"),
        ("5.", "Manufacturer Dashboard"),
        ("  5.1", "Dashboard Overview"),
        ("  5.2", "Statistics Cards"),
        ("  5.3", "Recent Schools"),
        ("6.", "Schools Management"),
        ("  6.1", "Schools List View"),
        ("  6.2", "Adding a New School"),
        ("  6.3", "Search & Filter"),
        ("  6.4", "Deleting a School"),
        ("7.", "School Detail Page"),
        ("  7.1", "Overview Tab"),
        ("  7.2", "Main Teacher Login Credentials"),
        ("  7.3", "School Logo Upload"),
        ("  7.4", "Classes Tab"),
        ("  7.5", "Students Tab"),
        ("  7.6", "Template Tab (JPG Template Mapper)"),
        ("  7.7", "Generate Tab (Batch Generation)"),
        ("  7.8", "Batches Tab"),
        ("  7.9", "Export Tab"),
        ("  7.10", "Bulk Import & Bulk Photo Upload"),
        ("8.", "Teacher Dashboard"),
        ("  8.1", "Overview Tab"),
        ("  8.2", "Class Form Links"),
        ("  8.3", "Class Breakdown & Export"),
        ("  8.4", "Students Tab"),
        ("  8.5", "Sub-Teachers Management"),
        ("  8.6", "ID Template Preview"),
        ("9.", "Student Submission Form"),
        ("  9.1", "Step 1: Fill Details"),
        ("  9.2", "Step 2: Photo Upload & Verification"),
        ("  9.3", "Step 3: Review & Submit"),
        ("  9.4", "Success Screen"),
        ("10.", "Photo Verification Engine"),
        ("11.", "ID Card Preview & Watermark"),
        ("12.", "Data Export (CSV & Excel)"),
        ("13.", "Security Features"),
        ("14.", "Troubleshooting & FAQ"),
    ]
    
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        if num.strip().endswith('.'):
            run = p.add_run(f"{num} {title_text}")
            run.bold = True
            run.font.size = Pt(12)
        else:
            run = p.add_run(f"    {num} {title_text}")
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 1. INTRODUCTION
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '1. Introduction & Overview', level=1)
    
    doc.add_paragraph(
        'Print ID Craft is a comprehensive, web-based platform designed for managing student ID cards '
        'across multiple schools. It provides a complete end-to-end solution from student data collection '
        'to ID card printing, with built-in role-based access control, photo verification, and batch processing.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Key Features', level=2)
    
    features = [
        ('Multi-School Management', 'Manage unlimited schools from a single manufacturer dashboard with dedicated credentials per school.'),
        ('Role-Based Access', 'Three distinct roles — Manufacturer (Admin), Main Teacher, and Sub-Teacher — each with appropriate permissions.'),
        ('Smart Form Links', 'Generate unique submission links per class that students/parents use to submit ID card data and photos.'),
        ('Photo Verification Engine', 'AI-powered photo quality checks ensuring passport-size format, proper background, and minimum resolution.'),
        ('JPG Template Mapping', 'Upload actual ID card template images and map data fields visually to precise positions on the card.'),
        ('Dual-Side Support', 'Full support for both front and back sides of ID cards in templates and batch generation.'),
        ('Batch Generation', 'Generate print-ready PDF batches with all approved student cards for efficient printing.'),
        ('Data Export', 'Export student data to CSV or Excel format with class and status filters.'),
        ('Bulk Import', 'Import student data from CSV/Excel files with validation and preview before committing.'),
        ('Bulk Photo Upload', 'Upload entire folders of photos with automatic name-matching to student records.'),
        ('Sub-Teacher System', 'Main teachers can create sub-teachers assigned to specific classes with restricted visibility.'),
        ('Watermarked Previews', 'All ID card previews display mandatory watermarks to prevent unauthorized reproduction.'),
    ]
    
    for title_text, desc_text in features:
        p = doc.add_paragraph()
        run = p.add_run(f'✅ {title_text}: ')
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x15, 0x80, 0x3D)
        run2 = p.add_run(desc_text)
        run2.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '2. System Architecture & User Roles', level=1)
    
    doc.add_paragraph(
        'Print ID Craft uses a modern web architecture built with Next.js, Prisma ORM, PostgreSQL (Supabase), '
        'and Supabase Storage for file uploads. The application supports three distinct user roles:'
    )
    
    doc.add_paragraph()
    
    # Roles table
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Role', 'Access Level', 'Key Permissions', 'Login URL']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        for p in table.cell(0, i).paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    
    roles = [
        ('Manufacturer\n(Admin)', 'Full System', 'Create schools, manage templates,\nview all data, generate batches,\nbulk download cards', '/login?mode=admin'),
        ('Main Teacher', 'School-wide', 'View all classes, approve/flag students,\nadd sub-teachers, add classes,\nexport data, map templates', '/login'),
        ('Sub-Teacher\n(Class Teacher)', 'Single Class', 'View assigned class only,\napprove/flag own students,\nadd comments', '/login'),
    ]
    
    for row_idx, (role, access, perms, url) in enumerate(roles, 1):
        table.cell(row_idx, 0).text = role
        table.cell(row_idx, 1).text = access
        table.cell(row_idx, 2).text = perms
        table.cell(row_idx, 3).text = url
        for col in range(4):
            for p in table.cell(row_idx, col).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    
    doc.add_paragraph()
    
    add_info_box(doc, "Workflow Summary",
        "Manufacturer creates school → Assigns Main Teacher → Main Teacher adds classes & sub-teachers → "
        "Form links shared with students → Students submit data & photos → Teachers review & approve → "
        "Manufacturer generates print batches → Cards printed & delivered.")
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 3. LANDING PAGE
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '3. Landing Page', level=1)
    
    doc.add_paragraph(
        'The landing page is the first thing visitors see when they access Print ID Craft. '
        'It features a modern, animated design with a clean layout that introduces the platform\'s capabilities.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '3.1 Hero Section', level=2)
    doc.add_paragraph(
        'The hero section displays the "We build powerful ID Card systems" tagline with call-to-action buttons '
        'for exploring the solution and logging in. On the right side, information cards highlight the platform\'s '
        'key differentiators.'
    )
    add_screenshot(doc, "01_landing_hero.png", "Figure 3.1: Landing Page — Hero Section with Navigation Bar")
    
    doc.add_paragraph()
    add_heading_styled(doc, '3.2 Statistics Bar', level=2)
    doc.add_paragraph(
        'An animated statistics bar displays key metrics including schools managed, cards printed, '
        'and match accuracy. Numbers animate from zero to their final values as the user scrolls.'
    )
    add_screenshot(doc, "03_landing_stats.png", "Figure 3.2: Animated Statistics Bar")
    
    doc.add_paragraph()
    add_heading_styled(doc, '3.3 Partnership Section', level=2)
    doc.add_paragraph(
        'The partnership section addresses common concerns schools face with traditional ID card vendors '
        'and positions Print ID Craft as a reliable solution. An animated submarine illustration adds '
        'visual interest.'
    )
    add_screenshot(doc, "04_landing_partnership.png", "Figure 3.3: Partnership-First Approach Section")
    
    doc.add_paragraph()
    add_heading_styled(doc, '3.4 Footer & Call-to-Action', level=2)
    doc.add_paragraph(
        'The dark-themed footer reinforces the value proposition and includes a prominent "Get Started" '
        'button, along with links to Privacy Policy, Terms of Service, and Support.'
    )
    add_screenshot(doc, "05_landing_footer.png", "Figure 3.4: Footer with Call-to-Action")
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 4. LOGIN SYSTEM
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '4. Login System', level=1)
    
    doc.add_paragraph(
        'Print ID Craft provides separate login portals for Teachers and Manufacturers (Admins). '
        'The login system uses NextAuth.js with secure credential-based authentication and bcrypt password hashing.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '4.1 Teacher Login', level=2)
    doc.add_paragraph(
        'The teacher login portal (accessible at /login) is designed for school administrators (Main Teachers) '
        'and class teachers (Sub-Teachers). It features a split-screen layout with an animated ID card mock-up '
        'on the left and the login form on the right.'
    )
    add_screenshot(doc, "06_teacher_login.png", "Figure 4.1: Teacher Login Page")
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Teacher Login Features:', level=3)
    features_list = [
        'Email and password authentication',
        'Photo guidelines reminder for student submissions',
        'Animated ID card illustration',
        'Responsive design for mobile devices',
        'Loading spinner during authentication',
        'Toast notifications for success/error feedback',
    ]
    for f in features_list:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_paragraph()
    add_heading_styled(doc, '4.2 Manufacturer (Admin) Login', level=2)
    doc.add_paragraph(
        'The manufacturer login (accessible at /login?mode=admin) is specifically for the ID card '
        'manufacturer/platform administrator. It displays a distinct "Manufacturer Login" header '
        'with "Manage schools, templates & printing" subtitle.'
    )
    add_screenshot(doc, "07_manufacturer_login.png", "Figure 4.2: Manufacturer (Admin) Login Page")
    
    doc.add_paragraph()
    add_step(doc, 1, 'Navigate to the login page (/login for teachers, /login?mode=admin for manufacturer)')
    add_step(doc, 2, 'Enter your registered email address')
    add_step(doc, 3, 'Enter your password')
    add_step(doc, 4, 'Click "Sign In" button')
    add_step(doc, 5, 'You will be redirected to the appropriate dashboard based on your role')
    
    add_info_box(doc, "Default Credentials",
        "When a school is created, the main teacher account is auto-generated with the default password: Teacher@123. "
        "The manufacturer should share these credentials with the school administrator.", "tip")
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 5. MANUFACTURER DASHBOARD
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '5. Manufacturer Dashboard', level=1)
    
    doc.add_paragraph(
        'After successful manufacturer login, the admin is redirected to the Manufacturer Dashboard. '
        'This is the central command center for managing all schools, templates, and printing operations.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '5.1 Dashboard Overview', level=2)
    doc.add_paragraph(
        'The manufacturer dashboard provides a high-level overview of the platform\'s activity. '
        'It includes a sidebar navigation with links to Dashboard and Schools pages.'
    )

    doc.add_paragraph()
    add_heading_styled(doc, '5.2 Statistics Cards', level=2)
    doc.add_paragraph(
        'Four animated statistics cards display key metrics at a glance:'
    )
    
    stats = [
        ('Total Schools', 'Number of registered schools in the system', 'Blue'),
        ('Total Students', 'Total student submissions across all schools', 'Purple'),
        ('Print Batches', 'Total number of generated print batches', 'Amber'),
        ('This Month', 'Student submissions for the current month', 'Green'),
    ]
    
    table = doc.add_table(rows=len(stats)+1, cols=3)
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Metric', 'Description', 'Color']):
        table.cell(0, i).text = h
        for r in table.cell(0, i).paragraphs[0].runs:
            r.bold = True
    for idx, (metric, description, color) in enumerate(stats, 1):
        table.cell(idx, 0).text = metric
        table.cell(idx, 1).text = description
        table.cell(idx, 2).text = color
    
    doc.add_paragraph()
    add_heading_styled(doc, '5.3 Recent Schools', level=2)
    doc.add_paragraph(
        'Below the statistics, a table displays the 5 most recently added schools with columns for '
        'School name, Student count, Classes count, Template status, and a Manage action button. '
        'On mobile devices, this table transforms into stacked cards for better readability.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Sidebar Navigation', level=3)
    doc.add_paragraph(
        'The left sidebar includes:'
    )
    nav_items = [
        'Dashboard — Overview with stats and recent schools',
        'Schools — Complete school list with management options',
        'User Profile — Shows logged-in user name with sign-out option',
    ]
    for item in nav_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 6. SCHOOLS MANAGEMENT
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '6. Schools Management', level=1)
    
    doc.add_paragraph(
        'The Schools page is accessible from the sidebar and provides complete management of all registered schools.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '6.1 Schools List View', level=2)
    doc.add_paragraph(
        'Schools are displayed as visually appealing cards with colored gradient banners. Each card shows:'
    )
    card_info = [
        'School name and initial avatar',
        'Contact email or address',
        'Student count and class count',
        'Template status badge (Ready/Not Set)',
        'Manage and Delete action buttons',
    ]
    for item in card_info:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph(
        'At the top, summary statistics show Total Schools, Total Students, and Total Classes across the platform.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '6.2 Adding a New School', level=2)
    doc.add_paragraph('To add a new school:')
    add_step(doc, 1, 'Click the "+ Add School" button at the top-right corner')
    add_step(doc, 2, 'Fill in the required fields: School Name and Contact Email')
    add_step(doc, 3, 'Optionally add the school Address and upload a School Logo')
    add_step(doc, 4, 'Add class names (e.g., "1st A", "2nd B", "3rd C") — you can add more later')
    add_step(doc, 5, 'Click "Create School" to save')
    
    doc.add_paragraph(
        'Upon creation, the system automatically generates a main teacher account for the school '
        'and creates all specified classes with unique form submission links.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '6.3 Search & Filter', level=2)
    doc.add_paragraph(
        'A search bar at the top allows you to search schools by name, email, or address. '
        'Search is debounced (400ms delay) for smooth performance. A clear button resets the filter. '
        'Pagination controls appear when there are more than 50 schools.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '6.4 Deleting a School', level=2)
    doc.add_paragraph(
        'To delete a school, click the red trash icon on the school card. A confirmation prompt '
        'requires you to type "DELETE" to confirm. This action permanently removes the school and ALL '
        'associated data (classes, students, templates, batches).'
    )
    
    add_info_box(doc, "Warning",
        "Deleting a school is irreversible. All student data, photos, classes, and generated batches will be permanently lost.",
        "warning")
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 7. SCHOOL DETAIL PAGE
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '7. School Detail Page', level=1)
    
    doc.add_paragraph(
        'Clicking "Manage" on any school card opens the detailed school management page. '
        'This page has 7 tabs: Overview, Classes, Students, Template, Generate, Batches, and Export.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '7.1 Overview Tab', level=2)
    doc.add_paragraph(
        'The Overview tab shows a quick summary of the school with 4 statistics cards:'
    )
    overview_stats = [
        'Total Classes — Number of classes/sections',
        'Total Students — Number of student submissions',
        'Print Batches — Number of generated print batches',
        'Template Status — Shows "Active" or "None"',
    ]
    for item in overview_stats:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    add_heading_styled(doc, '7.2 Main Teacher Login Credentials', level=2)
    doc.add_paragraph(
        'The Overview tab displays the Main Teacher Login section, which shows the credentials '
        'for the school\'s administrator account. This includes:'
    )
    cred_items = [
        'Login URL (the /login page)',
        'Email address (auto-generated from school email or manually set)',
        'Default password: Teacher@123',
        'Copy buttons for easy sharing',
        'Reset Password option to restore default password',
    ]
    for item in cred_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph(
        'If no teacher account exists, the manufacturer can either auto-generate one or manually enter an email.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '7.3 School Logo Upload', level=2)
    doc.add_paragraph(
        'Below the teacher credentials, a Logo Upload section allows the manufacturer to upload or replace '
        'the school logo. This logo appears on the school\'s ID cards. Supported formats: JPEG, PNG, WebP. '
        'Drag-and-drop upload is supported.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '7.4 Classes Tab', level=2)
    doc.add_paragraph(
        'The Classes tab provides complete class management with the following features:'
    )
    class_features = [
        'Add New Class — Enter name and optional expiry date',
        'Form Submission Links — Each class gets a unique URL that students use to submit their data',
        'Copy Link — One-click copy of the form link',
        'Share via WhatsApp — Pre-formatted message with school name and class details',
        'Share via Email — Opens email client with pre-filled subject and body',
        'Toggle Active/Inactive — Temporarily disable a class link to stop new submissions',
        'Set Expiry Date — Auto-deactivate links after a specified date',
        'Delete Class — Remove a class and all its student data (requires "DELETE" confirmation)',
    ]
    for item in class_features:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    add_heading_styled(doc, '7.5 Students Tab', level=2)
    doc.add_paragraph(
        'The Students tab displays all student submissions with powerful filtering and management tools:'
    )
    student_features = [
        'Filter by Status — All, Submitted, Approved, Flagged, Printed',
        'Filter by Class — Select a specific class to view its students',
        'Search by Name — Real-time search with debounced input',
        'Student Table — Photo, Serial Number, Name, Class, Status, Flag Note, Actions',
        'View Detail — Opens a full student detail modal with photo, form data, and ID card preview',
        'Approve — Mark a student as approved for printing',
        'Flag — Mark a student as flagged with a reason note',
        'Unflag — Remove the flag from a flagged student',
        'Bulk Import — Import students from CSV/Excel file',
        'Bulk Photo Upload — Upload entire photo folders with auto-matching',
        'Pagination — Navigate through pages of 50 students each',
    ]
    for item in student_features:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    add_heading_styled(doc, '7.6 Template Tab (JPG Template Mapper)', level=2)
    doc.add_paragraph(
        'The Template tab is one of the most powerful features. It allows the manufacturer to upload '
        'an actual JPG/PNG image of the ID card template and visually map where each data field should appear.'
    )
    doc.add_paragraph('How to set up a template:')
    add_step(doc, 1, 'Upload the ID card template image (front side)')
    add_step(doc, 2, 'Optionally upload a back side template')
    add_step(doc, 3, 'Use the visual mapper to add field zones — click or drag on the template image')
    add_step(doc, 4, 'For each zone, select the data field it maps to (e.g., Full Name, Class, Serial Number, Photo)')
    add_step(doc, 5, 'Adjust font size, color, alignment, and exact position for each mapping')
    add_step(doc, 6, 'Preview the card with sample data to verify placement')
    add_step(doc, 7, 'Save the template configuration')
    
    doc.add_paragraph()
    add_heading_styled(doc, '7.7 Generate Tab (Batch Generation)', level=2)
    doc.add_paragraph(
        'The Generate tab displays a preview of how ID cards will look with actual student data. '
        'The batch generator creates print-ready cards for all approved students.'
    )
    doc.add_paragraph('Generation process:')
    add_step(doc, 1, 'Review the card preview with actual student photos and data')
    add_step(doc, 2, 'Click "Generate Batch" to start processing')
    add_step(doc, 3, 'The system processes each approved student and generates high-resolution card images')
    add_step(doc, 4, 'A progress indicator shows the generation status')
    add_step(doc, 5, 'Once complete, the batch appears in the Batches tab for download')
    
    add_info_box(doc, "Note",
        "Only students with status 'SUBMITTED' or 'APPROVED' are included in batch generation. "
        "Flagged students are excluded. High-quality batch downloads are restricted to the Manufacturer role.",
        "note")
    
    doc.add_paragraph()
    add_heading_styled(doc, '7.8 Batches Tab', level=2)
    doc.add_paragraph(
        'The Batches tab lists all previously generated print batches with details including:'
    )
    batch_info = [
        'Batch creation date',
        'Number of students included',
        'Status (Processing/Ready)',
        'Download options: Front PDF, Back PDF, Manifest (CSV with student-card mapping)',
    ]
    for item in batch_info:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    add_heading_styled(doc, '7.9 Export Tab', level=2)
    doc.add_paragraph(
        'The Export tab allows downloading student data in structured formats:'
    )
    export_info = [
        'CSV Export — Comma-separated values for spreadsheet use',
        'Excel Export — Formatted .xlsx file with proper column headers',
        'Filter by Class — Export only specific class data',
        'Filter by Status — Export only students with a specific approval status',
    ]
    for item in export_info:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    add_heading_styled(doc, '7.10 Bulk Import & Bulk Photo Upload', level=2)
    
    doc.add_paragraph('Bulk Student Import:')
    add_step(doc, 1, 'Click "📥 Import" in the Students tab')
    add_step(doc, 2, 'Select a class for the imported students')
    add_step(doc, 3, 'Upload a CSV or Excel file with student data')
    add_step(doc, 4, 'Review the validation preview (shows matched columns, errors, warnings)')
    add_step(doc, 5, 'Confirm to import the data into the system')
    
    doc.add_paragraph()
    doc.add_paragraph('Bulk Photo Upload:')
    add_step(doc, 1, 'Click "📷 Bulk Photos" in the Students tab')
    add_step(doc, 2, 'Select an entire folder containing student photos')
    add_step(doc, 3, 'Photos are automatically matched to students by filename (serial number or student name)')
    add_step(doc, 4, 'Review results: Matched, Unmatched, and Error counts')
    add_step(doc, 5, 'For unmatched photos, use the Manual Assign feature to link them to specific students')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 8. TEACHER DASHBOARD
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '8. Teacher Dashboard', level=1)
    
    doc.add_paragraph(
        'The Teacher Dashboard is the central hub for school administrators and class teachers. '
        'It provides tools for managing students, sharing form links, and reviewing submissions.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '8.1 Overview Tab', level=2)
    doc.add_paragraph(
        'The overview tab shows:'
    )
    teacher_overview = [
        'School name and teacher role badge (Main Teacher or Class Teacher)',
        'Statistics cards: Total Submissions, Approved, Pending Review, Flagged',
        'Approval Progress bar showing percentage of approved students',
        'Class Form Links section with shareable URLs',
        'Class Breakdown table with student counts per class',
    ]
    for item in teacher_overview:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    add_heading_styled(doc, '8.2 Class Form Links', level=2)
    doc.add_paragraph(
        'For each class, a unique form link is displayed that can be shared with students/parents. '
        'Options include:'
    )
    link_features = [
        '📋 Copy — Copy the link to clipboard',
        '💬 WhatsApp — Share via WhatsApp with pre-formatted message including school name and class',
        'Class student count — Shows how many students have submitted for each class',
        'Assigned teacher name (if a sub-teacher is assigned)',
    ]
    for item in link_features:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph(
        'Main teachers can also add new classes directly from this section using the "+ Add Class" form.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '8.3 Class Breakdown & Export', level=2)
    doc.add_paragraph(
        'A detailed table shows each class with columns for Class Name, Assigned Teacher, Total Students, '
        'Approved count, and Flagged count. Export buttons allow downloading data in CSV or Excel format.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '8.4 Students Tab', level=2)
    doc.add_paragraph(
        'The Students tab provides a comprehensive table of all student submissions with columns for:'
    )
    student_cols = [
        'Photo — Thumbnail of the student\'s uploaded photo',
        'Serial — Unique serial number assigned to each submission',
        'Name — Student\'s full name from form data',
        'Class — The class the student belongs to',
        'Status — Color-coded badge (Submitted/Approved/Flagged/Printed)',
        'Comment — Teacher comments visible to the manufacturer',
        'Actions — View (👁), Approve (✓), Disapprove (✕), Flag (🚩), Edit (✏️), Comment (💬)',
    ]
    for item in student_cols:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Available actions per student:')
    actions_list = [
        ('View (👁)', 'Opens detailed student modal with photo, all form data, and ID card preview'),
        ('Approve (✓)', 'Marks the student as approved — eligible for printing'),
        ('Disapprove (✕)', 'Marks the student as flagged/disapproved'),
        ('Flag (🚩)', 'Flags the student with a custom reason note'),
        ('Unflag', 'Removes the flag and returns to submitted status'),
        ('Edit (✏️)', 'Opens a modal to edit all form fields for the student'),
        ('Comment (💬)', 'Adds a teacher comment visible to the manufacturer'),
    ]
    for action, description in actions_list:
        p = doc.add_paragraph()
        run = p.add_run(f'{action}: ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(description)
        run2.font.size = Pt(10)
    
    doc.add_paragraph()
    add_heading_styled(doc, '8.5 Sub-Teachers Management', level=2)
    doc.add_paragraph(
        'The Sub-Teachers tab (available only to Main Teachers) provides tools for managing class-specific teachers.'
    )
    doc.add_paragraph('Adding a Sub-Teacher:')
    add_step(doc, 1, 'Click "+ Add Teacher" button')
    add_step(doc, 2, 'Fill in: Name, Email, Password (min 6 characters), and Assign to Class')
    add_step(doc, 3, 'Click "Add Class Teacher"')
    add_step(doc, 4, 'A success banner appears showing the credentials (Name, Email, Password, Login URL)')
    add_step(doc, 5, 'Copy and share these credentials with the teacher — the password will NOT be shown again')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The sub-teachers list table shows Name, Email, Assigned Class, Added date, and a Remove action. '
        'Sub-teachers can only see and manage students in their assigned class.'
    )
    
    add_info_box(doc, "Security Note",
        "Sub-teacher passwords are stored securely using bcrypt hashing. The password is displayed only once "
        "at creation time. If forgotten, the main teacher must delete and recreate the sub-teacher account.",
        "warning")
    
    doc.add_paragraph()
    add_heading_styled(doc, '8.6 ID Template Preview', level=2)
    doc.add_paragraph(
        'The Template tab (available only to Main Teachers) shows a preview of the school\'s ID card template. '
        'The main teacher can use the JPG Template Mapper to visually position data fields on the card. '
        'All previews include mandatory watermarks to prevent unauthorized reproduction.'
    )
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 9. STUDENT SUBMISSION FORM
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '9. Student Submission Form', level=1)
    
    doc.add_paragraph(
        'Students (or parents) use a unique class-specific URL to submit their ID card details. '
        'The form follows a 3-step wizard process: Details → Photo → Review.'
    )
    
    add_screenshot(doc, "30_submit_error.png", "Figure 9.0: Invalid Form Link Error Page")
    
    doc.add_paragraph()
    add_heading_styled(doc, '9.1 Step 1: Fill Details', level=2)
    doc.add_paragraph(
        'The form displays the school name, class name, and all configured fields (e.g., Full Name, '
        'Father\'s Name, Date of Birth, Blood Group, Address, etc.). Required fields are marked with a '
        'red asterisk. Field types include text, tel, date, select, and textarea.'
    )
    doc.add_paragraph(
        'A step indicator at the top shows the current position: 1-Details, 2-Photo, 3-Review.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, '9.2 Step 2: Photo Upload & Verification', level=2)
    doc.add_paragraph(
        'After filling in details, the student uploads a passport-size photo. The system includes:'
    )
    photo_features = [
        'Photo Requirements Box — Lists all requirements (300px minimum, plain background, face visible, 3:4 ratio)',
        'Photo Verifier Component — AI-powered verification that checks background, face detection, and dimensions',
        'Sample Photo Reference — Shows the expected format visually',
        'Image Crop Tool — React Image Crop component for precise passport-size cropping',
        'File validation — Accepts JPEG, PNG, WebP; max 5MB; min 300×300 pixels',
    ]
    for item in photo_features:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    add_heading_styled(doc, '9.3 Step 3: Review & Submit', level=2)
    doc.add_paragraph(
        'The review page shows a side-by-side layout with:'
    )
    review_features = [
        'Card Preview — Shows how the ID card will look with the student\'s actual data and photo',
        'Details Check — Lists all entered fields for verification',
        'Student Photo Thumbnail — Shows the cropped photo',
        'Confirmation Note — Reminds that changes cannot be made after submission',
        'Progress Bar — Shows upload progress during submission (0-100%)',
        'Submit Registration button — Finalizes the submission',
    ]
    for item in review_features:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    add_heading_styled(doc, '9.4 Success Screen', level=2)
    doc.add_paragraph(
        'After successful submission, the student sees:'
    )
    success_features = [
        'Green checkmark confirmation',
        'Unique Serial Number — Students should save this for their records',
        'ID Card Preview — Shows the finalized card with a watermark',
        '"Powered by Print ID Craft" branding',
    ]
    for item in success_features:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 10. PHOTO VERIFICATION ENGINE
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '10. Photo Verification Engine', level=1)
    
    doc.add_paragraph(
        'Print ID Craft includes a built-in Photo Verification Engine that automatically validates '
        'uploaded photos against ID card standards before accepting them.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Verification Checks:', level=2)
    
    checks = [
        ('Minimum Resolution', 'Photos must be at least 300×300 pixels to ensure print quality'),
        ('File Format', 'Only JPEG, PNG, and WebP formats are accepted'),
        ('File Size', 'Maximum 5MB to ensure reasonable upload times'),
        ('Aspect Ratio', 'Photos are cropped to 3:4 ratio (passport-size standard)'),
        ('Background Check', 'Verifies that the photo has a plain/solid background'),
        ('Face Detection', 'Ensures a face is clearly visible and properly centered'),
    ]
    
    for title_text, desc_text in checks:
        p = doc.add_paragraph()
        run = p.add_run(f'🔍 {title_text}: ')
        run.bold = True
        run2 = p.add_run(desc_text)
    
    doc.add_paragraph()
    add_info_box(doc, "Photo Tips",
        "For best results: Use natural lighting, stand against a plain white/light wall, "
        "face the camera directly, and avoid filters or heavy editing.", "tip")
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 11. ID CARD PREVIEW & WATERMARK
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '11. ID Card Preview & Watermark', level=1)
    
    doc.add_paragraph(
        'All ID card previews throughout the platform display mandatory watermarks to prevent unauthorized '
        'reproduction. The system supports two types of card preview:'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, 'JPG Card Preview', level=2)
    doc.add_paragraph(
        'When a JPG template is configured with field mappings, the JpgCardPreview component renders '
        'the actual template image with student data overlaid at the mapped positions. This provides '
        'a WYSIWYG (What You See Is What You Get) preview of the final printed card.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Canvas-Based Preview', level=2)
    doc.add_paragraph(
        'For templates using the drag-and-drop designer, a canvas-based preview renders the card layout '
        'with proper positioning of text, photos, logos, and QR codes.'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Watermark System', level=2)
    doc.add_paragraph(
        'A diagonal watermark text (e.g., "Wise Melon") is overlaid on all card previews. This watermark:'
    )
    watermark_features = [
        'Appears on all preview screens (student submission, teacher dashboard, manufacturer view)',
        'Is removed only in the high-quality batch downloads (restricted to Manufacturer role)',
        'Cannot be disabled by teachers or students',
        'Covers the entire card surface to prevent screenshot-based reproduction',
    ]
    for item in watermark_features:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 12. DATA EXPORT
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '12. Data Export (CSV & Excel)', level=1)
    
    doc.add_paragraph(
        'Both the manufacturer and teacher dashboards support exporting student data.'
    )
    
    doc.add_paragraph()
    export_table = doc.add_table(rows=3, cols=3)
    export_table.style = 'Light Grid Accent 1'
    headers = ['Feature', 'CSV Export', 'Excel Export']
    for i, h in enumerate(headers):
        export_table.cell(0, i).text = h
        for r in export_table.cell(0, i).paragraphs[0].runs:
            r.bold = True
    
    export_data = [
        ('Format', '.csv (comma-separated)', '.xlsx (formatted workbook)'),
        ('Filters', 'Class, Status', 'Class, Status'),
    ]
    for idx, (feature, csv_val, excel_val) in enumerate(export_data, 1):
        export_table.cell(idx, 0).text = feature
        export_table.cell(idx, 1).text = csv_val
        export_table.cell(idx, 2).text = excel_val
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Exported files include: Serial Number, Student Name, Class, all form fields, Photo URL, '
        'Status, Flag Note, Teacher Comment, and Submission Date.'
    )
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 13. SECURITY FEATURES
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '13. Security Features', level=1)
    
    security_items = [
        ('Authentication', 'NextAuth.js with JWT-based sessions (24-hour expiry). Passwords hashed with bcrypt.'),
        ('Role-Based Access Control', 'Middleware enforces role-based routing — Teachers cannot access manufacturer pages and vice versa.'),
        ('Sub-Teacher Isolation', 'Class teachers can only view and manage students in their assigned class.'),
        ('Watermarked Previews', 'All ID card previews have mandatory watermarks to prevent unauthorized use.'),
        ('Batch Download Restriction', 'High-quality, unwatermarked batch downloads are restricted to the Manufacturer role only.'),
        ('Form Link Control', 'Class form links can be deactivated or set with expiry dates to control submission windows.'),
        ('Delete Confirmation', 'Destructive operations (delete school, delete class) require typing "DELETE" to confirm.'),
        ('Input Validation', 'All form inputs are validated on both client and server side.'),
        ('Photo Validation', 'Photos are validated for format, size, resolution, and content quality.'),
        ('Force Logout', 'API endpoint available to force sign-out a specific user session.'),
    ]
    
    for title_text, desc_text in security_items:
        p = doc.add_paragraph()
        run = p.add_run(f'🔒 {title_text}: ')
        run.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        run2 = p.add_run(desc_text)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # 14. TROUBLESHOOTING
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, '14. Troubleshooting & FAQ', level=1)
    
    faqs = [
        ('Q: I forgot my teacher login password. What do I do?',
         'A: Contact the manufacturer/admin. They can reset your password from the School Detail → Overview tab using the "Reset Pw" button. The password will be reset to the default: Teacher@123.'),
        ('Q: A student submitted incorrect information. How do I fix it?',
         'A: Go to the Students tab, find the student, and click the Edit (✏️) button. You can modify any form field. Alternatively, flag the student and add a comment explaining the issue.'),
        ('Q: The form link isn\'t working for students.',
         'A: Check if the class link is active (Classes tab → Toggle). Also check if an expiry date has passed. Generate a new link if needed.'),
        ('Q: Photos are being rejected during upload.',
         'A: Ensure the photo meets all requirements: at least 300×300 pixels, JPEG/PNG/WebP format, under 5MB, plain background, and front-facing.'),
        ('Q: How do I generate print-ready ID cards?',
         'A: Navigate to the school\'s Generate tab, review the preview, and click "Generate Batch". The batch will appear in the Batches tab when ready.'),
        ('Q: Can sub-teachers see all students in the school?',
         'A: No. Sub-teachers can only view and manage students in their assigned class. Only the Main Teacher has school-wide visibility.'),
        ('Q: How do I share form links with parents?',
         'A: Go to the Teacher Dashboard → Overview → Class Form Links section. Use the Copy button or the WhatsApp share button for instant sharing.'),
        ('Q: Can I import existing student data from a spreadsheet?',
         'A: Yes. In the Students tab, click "📥 Import" and upload a CSV or Excel file. The system will validate the data before importing.'),
        ('Q: Why do ID card previews have a watermark?',
         'A: Watermarks prevent unauthorized reproduction. They appear on all preview screens but are removed in the final print-ready batch downloads.'),
        ('Q: How do I upload photos in bulk?',
         'A: In the Students tab, click "📷 Bulk Photos" and select an entire folder. Photos are automatically matched to students by filename. Unmatched photos can be manually assigned.'),
    ]
    
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        p2 = doc.add_paragraph()
        run2 = p2.add_run(a)
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        doc.add_paragraph()
    
    # ═══════════════════════════════════════════════
    # FINAL PAGE
    # ═══════════════════════════════════════════════
    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph()
    
    end_title = doc.add_heading('Thank You', level=1)
    end_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    end_text = doc.add_paragraph(
        'For additional support or feature requests, please contact the Print ID Craft team.'
    )
    end_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    end_url = doc.add_paragraph('www.printidcraft.com')
    end_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_url.runs[0].font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    end_url.runs[0].font.size = Pt(14)
    end_url.runs[0].bold = True
    
    doc.add_paragraph()
    version_final = doc.add_paragraph('© 2026 Print ID Craft. All rights reserved.')
    version_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_final.runs[0].font.size = Pt(10)
    version_final.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    # Save the document
    doc.save(OUTPUT_PATH)
    print(f"\n✅ User Manual saved to: {OUTPUT_PATH}")
    print(f"📄 Document created with screenshots from: {SCREENSHOTS_DIR}")

if __name__ == "__main__":
    create_manual()
